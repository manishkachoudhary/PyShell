import os
import subprocess
from colorama import init, Fore, Style
from fpdf import FPDF
import fitz  # PyMuPDF
from docx import Document


try:
    import readline  # for arrow-key history support (Linux/macOS)
except ImportError:
    try:
        import pyreadline3 as readline  # for Windows
    except ImportError:
        readline = None

init(autoreset=True)

# ------------------------
# Data Structures
# ------------------------
class Stack:
    def __init__(self): self.items = []
    def push(self, item): self.items.append(item)
    def pop(self): return self.items.pop() if self.items else None
    def peek(self): return self.items[-1] if self.items else None
    def display(self): return self.items
    def trace(self): return self.items[::-1]
    def sort(self, reverse=False): self.items.sort(key=int, reverse=reverse)
    def search(self, target):
        try: return f"{target} found at position {self.items[::-1].index(target)+1} from top"
        except ValueError: return f"{target} not found"
    def bsearch(self, val):
        try:
            sorted_items = sorted(self.items, key=int)
            return f"{val} found at index {sorted_items.index(val)} in sorted stack"
        except ValueError:
            return f"{val} not found in sorted stack"

class Queue:
    def __init__(self): self.items = []
    def enqueue(self, item): self.items.append(item)
    def dequeue(self): return self.items.pop(0) if self.items else None
    def front(self): return self.items[0] if self.items else None
    def rear(self): return self.items[-1] if self.items else None
    def display(self): return self.items
    def sort(self, reverse=False): self.items.sort(key=int, reverse=reverse)
    def search(self, target):
        try: return f"{target} found at position {self.items.index(target)+1} from front"
        except ValueError: return f"{target} not found"
    def bsearch(self, val):
        try:
            sorted_items = sorted(self.items, key=int)
            return f"{val} found at index {sorted_items.index(val)} in sorted queue"
        except ValueError:
            return f"{val} not found in sorted queue"

class Node:
    def __init__(self, data): self.data = data; self.next = None

class LinkedList:
    def __init__(self): self.head = None
    def insert(self, data):
        node = Node(data)
        if not self.head: self.head = node
        else:
            cur = self.head
            while cur.next: cur = cur.next
            cur.next = node
    def delete(self, data):
        cur, prev = self.head, None
        while cur:
            if cur.data == data:
                if prev: prev.next = cur.next
                else: self.head = cur.next
                return True
            prev, cur = cur, cur.next
        return False
    def display(self):
        result, cur = [], self.head
        while cur: result.append(cur.data); cur = cur.next
        return result
    def reverse_display(self):
        def recurse(node): return recurse(node.next) + [node.data] if node else []
        return recurse(self.head)
    def sort(self, reverse=False):
        values = self.display()
        values.sort(key=int, reverse=reverse)
        self.head = None
        for v in values: self.insert(v)
    def search(self, target):
        cur, idx = self.head, 1
        while cur:
            if cur.data == target: return f"{target} found at position {idx}"
            cur = cur.next; idx += 1
        return f"{target} not found"
    def bsearch(self, val):
        try:
            values = sorted(self.display(), key=int)
            return f"{val} found at index {values.index(val)} in sorted list"
        except ValueError:
            return f"{val} not found in sorted list"

class CircularLinkedList:
    def __init__(self): self.tail = None
    def insert(self, data):
        node = Node(data)
        if not self.tail:
            self.tail = node
            node.next = node
        else:
            node.next = self.tail.next
            self.tail.next = node
            self.tail = node
    def delete(self, data):
        if not self.tail: return False
        cur, prev = self.tail.next, self.tail
        while True:
            if cur.data == data:
                if cur == self.tail and cur.next == self.tail:
                    self.tail = None
                else:
                    prev.next = cur.next
                    if cur == self.tail:
                        self.tail = prev
                return True
            prev, cur = cur, cur.next
            if cur == self.tail.next: break
        return False
    def display(self):
        result = []
        if not self.tail: return result
        cur = self.tail.next
        while True:
            result.append(cur.data)
            cur = cur.next
            if cur == self.tail.next: break
        return result
    def sort(self, reverse=False):
        items = self.display()
        items.sort(key=int, reverse=reverse)
        self.tail = None
        for val in items: self.insert(val)
    def search(self, target):
        cur = self.tail.next if self.tail else None
        idx = 1
        while cur:
            if cur.data == target:
                return f"{target} found at position {idx}"
            cur = cur.next
            idx += 1
            if cur == self.tail.next: break
        return f"{target} not found"
    def bsearch(self, val):
        try:
            values = sorted(self.display(), key=int)
            return f"{val} found at index {values.index(val)} in sorted clist"
        except ValueError:
            return f"{val} not found in sorted clist"

class DoublyNode:
    def __init__(self, data): self.data = data; self.prev = None; self.next = None

class DoublyLinkedList:
    def __init__(self): self.head = None
    def insert(self, data):
        node = DoublyNode(data)
        if not self.head: self.head = node
        else:
            cur = self.head
            while cur.next: cur = cur.next
            cur.next = node
            node.prev = cur
    def delete(self, data):
        cur = self.head
        while cur:
            if cur.data == data:
                if cur.prev: cur.prev.next = cur.next
                else: self.head = cur.next
                if cur.next: cur.next.prev = cur.prev
                return True
            cur = cur.next
        return False
    def display(self):
        result, cur = [], self.head
        while cur: result.append(cur.data); cur = cur.next
        return result
    def reverse_display(self):
        result = []
        cur = self.head
        while cur and cur.next: cur = cur.next
        while cur: result.append(cur.data); cur = cur.prev
        return result
    def sort(self, reverse=False):
        values = self.display()
        values.sort(key=int, reverse=reverse)
        self.head = None
        for v in values: self.insert(v)
    def search(self, target):
        cur, idx = self.head, 1
        while cur:
            if cur.data == target: return f"{target} found at position {idx}"
            cur = cur.next; idx += 1
        return f"{target} not found"
    def bsearch(self, val):
        try:
            values = sorted(self.display(), key=int)
            return f"{val} found at index {values.index(val)} in sorted dlist"
        except ValueError:
            return f"{val} not found in sorted dlist"

# ------------------------
# Core Shell + File + DSA Commands
# ------------------------
stack = Stack()
queue = Queue()
linkedlist = LinkedList()
clist = CircularLinkedList()
dlist = DoublyLinkedList()

def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

def execute_command(command):
    command = command.strip()
    if not command: return

    # I/O Redirection and Piping
    if '|' in command or '>' in command or '<' in command:
        try:
            if '|' in command:
                # Handle multi-stage piping
                cmds = [c.strip() for c in command.split('|')]
                prev_process = None
                for i, cmd in enumerate(cmds):
                    if i == 0:
                        prev_process = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                    else:
                        prev_process = subprocess.Popen(cmd, shell=True, stdin=prev_process.stdout, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                stdout, stderr = prev_process.communicate()
                if stdout:
                    print(Fore.GREEN + stdout, end='')
                if stderr:
                    print(Fore.RED + stderr, end='')
            else:
                # Single-stage redirection (no pipe)
                output = subprocess.check_output(command, shell=True, text=True, stderr=subprocess.STDOUT)
                print(Fore.GREEN + output)
        except subprocess.CalledProcessError as e:
            print(Fore.RED + e.output)
        except Exception as e:
            print(Fore.RED + f"Error: {e}")
        return

    args = command.split()
    cmd = args[0].lower()

    def result(val): print(Fore.GREEN + str(val))

    try:
        if cmd == 'cd':
            target = ' '.join(args[1:])
            if target.endswith(':') and os.path.exists(target + "\\"):
                os.chdir(target + "\\")
            else:
                os.chdir(target)
            print(Fore.GREEN + f"Changed directory to {os.getcwd()}")

        elif cmd == 'pwd':
            print(Fore.GREEN + os.getcwd())

        elif cmd == 'file':
            action = args[1]
            filename = ' '.join(args[2:])
            if action == 'read':
                with open(filename, 'r', encoding='utf-8') as f:
                    print(Fore.GREEN + f.read())
            elif action == 'write':
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(input("Enter text to write: "))
                    print(Fore.GREEN + "File written successfully.")
            elif action == 'append':
                with open(filename, 'a', encoding='utf-8') as f:
                    f.write(input("Enter text to append: "))
                    print(Fore.GREEN + "Text appended successfully.")
            elif action == 'delete':
                os.remove(filename)
                print(Fore.GREEN + "File deleted.")
            elif action == 'access':
                exists = os.access(filename, os.R_OK)
                print(Fore.GREEN + (f"{filename} is accessible." if exists else f"{filename} is not accessible."))
            elif action == 'sneha':
                os.startfile(filename)
            else:
                print(Fore.RED + "Unknown file operation.")
        
        elif cmd == 'pdf':
            if len(args) >= 3 and args[1] == 'write':
                filename = args[2]
                if not filename.lower().endswith('.pdf'):
                    print(Fore.RED + "Error: Filename must end with .pdf")
                    return

                print("Enter text to write in PDF (type ':q' on a new line to finish):")
                lines = []
                while True:
                    line = input()
                    if line.strip() == ':q':
                        break
                    lines.append(line)

                text = "\n".join(lines)
                try:
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, txt=text)
                    pdf.output(filename)
                    print(Fore.GREEN + "PDF written successfully.")
                except Exception as e:
                    print(Fore.RED + f"Failed to write PDF: {e}")

            elif len(args) >= 3 and args[1] == 'read':
                filename = args[2]
                if not filename.lower().endswith('.pdf'):
                    print(Fore.RED + "Error: Not a .pdf file")
                    return
                try:
                    import fitz  # pymupdf (put this at top of file too)
                    doc = fitz.open(filename)
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        text = page.get_text()
                        print(Fore.GREEN + f"\n--- Page {page_num + 1} ---\n{text}")
                    doc.close()
                except Exception as e:
                    print(Fore.RED + f"Error reading PDF: {e}")

            else:
                print(Fore.RED + "Usage: pdf write filename.pdf OR pdf read filename.pdf")

        elif cmd == 'docx':
            if len(args) >= 3 and args[1] == 'write':
                filename = args[2]
                if not filename.lower().endswith('.docx'):
                    print(Fore.RED + "Error: Filename must end with .docx")
                    return

                print("Enter text to write in DOCX (type ':q' on a new line to finish):")
                lines = []
                while True:
                    line = input()
                    if line.strip() == ':q':
                        break
                    lines.append(line)

                try:
                    doc = Document()
                    for line in lines:
                        doc.add_paragraph(line)
                    doc.save(filename)
                    print(Fore.GREEN + "DOCX file written successfully.")
                except Exception as e:
                    print(Fore.RED + f"Error writing DOCX: {e}")

            elif len(args) >= 3 and args[1] == 'read':
                filename = args[2]
                if not filename.lower().endswith('.docx'):
                    print(Fore.RED + "Error: Not a .docx file")
                    return

                try:
                    doc = Document(filename)
                    for para in doc.paragraphs:
                        if para.text.strip():
                            print(Fore.GREEN + para.text)
                except Exception as e:
                    print(Fore.RED + f"Error reading DOCX: {e}")

            else:
                print(Fore.RED + "Usage: docx write filename.docx OR docx read filename.docx")



        # DSA Commands
        elif cmd == "stack":
            sub = args[1]; val = args[2] if len(args) > 2 else ''
            if sub == "push": [stack.push(x) for x in args[2:]]
            elif sub == "pop": result(stack.pop())
            elif sub == "peek": result(stack.peek())
            elif sub == "trace": result(stack.trace())
            elif sub == "display": result(stack.display())
            elif sub == "sort": stack.sort('reverse' in args); result(stack.display())
            elif sub == "search": result(stack.search(val))
            elif sub == "bsearch": result(stack.bsearch(val))

        elif cmd == "queue":
            sub = args[1]; val = args[2] if len(args) > 2 else ''
            if sub == "enqueue": [queue.enqueue(x) for x in args[2:]]
            elif sub == "dequeue": result(queue.dequeue())
            elif sub == "front": result(queue.front())
            elif sub == "rear": result(queue.rear())
            elif sub == "display": result(queue.display())
            elif sub == "sort": queue.sort('reverse' in args); result(queue.display())
            elif sub == "search": result(queue.search(val))
            elif sub == "bsearch": result(queue.bsearch(val))

        elif cmd == "linkedlist":
            sub = args[1]; val = args[2] if len(args) > 2 else ''
            if sub == "insert": [linkedlist.insert(x) for x in args[2:]]
            elif sub == "delete": result(linkedlist.delete(val))
            elif sub == "display": result(linkedlist.display())
            elif sub == "reverse": result(linkedlist.reverse_display())
            elif sub == "sort": linkedlist.sort('reverse' in args); result(linkedlist.display())
            elif sub == "search": result(linkedlist.search(val))
            elif sub == "bsearch": result(linkedlist.bsearch(val))

        elif cmd == "clist":
            sub = args[1]; val = args[2] if len(args) > 2 else ''
            if sub == "insert": [clist.insert(x) for x in args[2:]]
            elif sub == "delete": result(clist.delete(val))
            elif sub == "display": result(clist.display())
            elif sub == "sort": clist.sort('reverse' in args); result(clist.display())
            elif sub == "search": result(clist.search(val))
            elif sub == "bsearch": result(clist.bsearch(val))

        elif cmd == "dlist":
            sub = args[1]; val = args[2] if len(args) > 2 else ''
            if sub == "insert": [dlist.insert(x) for x in args[2:]]
            elif sub == "delete": result(dlist.delete(val))
            elif sub == "display": result(dlist.display())
            elif sub == "reverse": result(dlist.reverse_display())
            elif sub == "sort": dlist.sort('reverse' in args); result(dlist.display())
            elif sub == "search": result(dlist.search(val))
            elif sub == "bsearch": result(dlist.bsearch(val))

        elif cmd in ['clear', 'cls']:
            clear_screen()

        elif cmd == 'help':
            print(Fore.CYAN + """
Available Commands:
DSA Commands: stack, queue, linkedlist, clist, dlist 
  (supports push, pop, insert, delete, display, sort, reverse, search, binary search)
File Ops: file read/write/append/delete/access/open filename
Navigation: cd path | pwd | clear | help | exit
System: any OS command (e.g., dir, type)
""")

        elif cmd == 'exit':
            print(Fore.YELLOW + "Exiting PyShell...")
            exit()

        else:
            result = subprocess.run(command, shell=True, capture_output=True, text=True)
            if result.stdout:
                print(Fore.GREEN + result.stdout)
            if result.stderr:
                print(Fore.RED + result.stderr)

    except Exception as e:
        print(Fore.RED + f"Error: {e}")

# ------------------------
# Main Loop
# ------------------------
def main():
    print(Fore.CYAN + "Welcome to PyShell: DSA + FileOps + System Commands")
    print(Fore.YELLOW + "Type 'help' for commands. Type 'exit' to quit.\n")
    while True:
        try:
            cmd = input(Fore.YELLOW + "PyShell> " + Style.RESET_ALL)
            execute_command(cmd)
        except (EOFError, KeyboardInterrupt):
            print(Fore.RED + "\nUse 'exit' to quit.")

if __name__ == "__main__":
    main()
